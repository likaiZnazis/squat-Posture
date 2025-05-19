from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import date

def create_report(total_records, sensor_freq, segment_length, form_counts,
                  accuracy, specificity, sensitivity, confusion_matrixPath):

    activeDirectoryName = "dataset"
    currentDirectory = os.getcwd()
    path = os.path.join(currentDirectory, activeDirectoryName)

    document = Document()

    # Title and intro
    title = document.add_heading('ATSKAITE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph('PIETUPIENA DATUKOPAI UN KLASIFIKĀCIJAI')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(str(date.today()), style='Normal')

    document.add_paragraph(
        '\tDokuments sastāv no divām daļām: datu kopa un modelis. Datu kopas daļā tiek aprakstītas dažādas specifikas par izveidoto datu kopu. '
        'Pretēji tas pats arī ir darīts modeļa nodaļā.\n'
    )

    # --- Section: Dataset ---
    document.add_heading('DATU KOPA', level=1)
    document.add_paragraph('Datu kopas, kopējās iezīmes')
    document.add_paragraph(f'Kopējais pietupienu skaits – {total_records}')
    document.add_paragraph(f'Sensora frekvence - {sensor_freq}')
    document.add_paragraph(f'Segmentu garums – {segment_length}')

    # Form counts table
    document.add_paragraph('Formu iedalījums ierakstu daudzumos:')
    table1 = document.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Formas nosaukums'
    hdr_cells[1].text = 'Ierakstu daudzums'
    for form, count in form_counts.items():
        row_cells = table1.add_row().cells
        row_cells[0].text = str(form)
        row_cells[1].text = str(count)

    document.add_page_break()
    # --- Section: Model Evaluation ---
    document.add_heading('MODELIS', level=1)
    document.add_paragraph(f'Precizitāte - {accuracy:.4f}')
    document.add_paragraph(f'Jutība - {sensitivity:.4f}')
    document.add_paragraph(f'Specifiskums - {specificity:.4f}')

    # Confusion Matrix Table
    # cm_table = document.add_table(rows=len(class_names)+1, cols=len(class_names)+1)
    # cm_table.style = 'Table Grid'

    # Header row
    # cm_table.cell(0, 0).text = " "
    # for j, name in enumerate(class_names):
    #     cm_table.cell(0, j+1).text = name
    #     cm_table.cell(j+1, 0).text = name

    # # Fill matrix
    # for i in range(len(confusion_matrix)):
    #     for j in range(len(confusion_matrix[i])):
    #         cm_table.cell(i+1, j+1).text = str(confusion_matrix[i][j])
    document.add_picture(confusion_matrixPath, width=Inches(5.5))
    # Save document
    document.save(os.path.join(path, "report.docx"))
